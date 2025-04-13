from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import json
import uuid
from datetime import datetime
import io
import markdown
from xhtml2pdf import pisa
from docx import Document
from openai import OpenAI
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics

import os
import json
from docx import Document


app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

#Paste your API key here
client = OpenAI(api_key="YOUR-API-KEY")

# Set the directory for storing user data
DATA_DIR = os.path.dirname(os.path.abspath(__file__))
USER_DATA_FILE = os.path.join(DATA_DIR, "user_data.json")
RESUMES_DIR = os.path.join(DATA_DIR, "resumes")
CHAT_HISTORY_FILE = os.path.join(DATA_DIR, "chat_history.json")
AI_CONVERSATIONS_FILE = os.path.join(DATA_DIR, "ai_conversations.json")

print(f"Data directory: {DATA_DIR}")
print(f"User data file path: {USER_DATA_FILE}")
print(f"Resumes directory: {RESUMES_DIR}")
print(f"Chat history file path: {CHAT_HISTORY_FILE}")
print(f"AI conversations file path: {AI_CONVERSATIONS_FILE}")

# Register Times New Roman font if available
try:
    pdfmetrics.registerFont(TTFont('Times-Roman', '/usr/share/fonts/truetype/msttcorefonts/times.ttf'))
except:
    pass  # fallback to built-in Times-Roman

# Create directories if they don't exist
if not os.path.exists(RESUMES_DIR):
    os.makedirs(RESUMES_DIR)
    print(f"Created resumes directory: {RESUMES_DIR}")

def load_user_data():
    if os.path.exists(USER_DATA_FILE):
        try:
            with open(USER_DATA_FILE, 'r') as f:
                data = json.load(f)
                print(f"Successfully loaded user data: {len(str(data))} bytes")
                return data
        except Exception as e:
            print(f"Error loading user data: {e}")
    
    # Default data if file doesn't exist or is corrupted
    default_data = {
        "name": "John Doe",
        "email": "john@example.com",
        "phone": "(123) 456-7890",
        "linkedin": "linkedin.com/in/johndoe",
        "education": [],
        "workExperiences": [],
        "skills": ["JavaScript", "React", "Python"],
        "resumes": [],
        "coverLetters": []
    }
    print("Creating default user data")
    return default_data

def save_user_data(data):
    try:
        print(f"Saving user data to {os.path.abspath(USER_DATA_FILE)}")
        with open(USER_DATA_FILE, 'w') as f:
            json.dump(data, f, indent=2)
        print(f"Data saved successfully. File size: {os.path.getsize(USER_DATA_FILE)} bytes")
        return True
    except Exception as e:
        print(f"Error saving user data: {e}")
        return False

def load_chat_history():
    if os.path.exists(CHAT_HISTORY_FILE):
        try:
            with open(CHAT_HISTORY_FILE, 'r') as f:
                data = json.load(f)
                print(f"Successfully loaded chat history: {len(data)} conversations")
                return data
        except Exception as e:
            print(f"Error loading chat history: {e}")
    
    # Default empty chat history if file doesn't exist or is corrupted
    return []

def save_chat_history(data):
    try:
        print(f"Saving chat history to {os.path.abspath(CHAT_HISTORY_FILE)}")
        with open(CHAT_HISTORY_FILE, 'w') as f:
            json.dump(data, f, indent=2)
        print(f"Chat history saved successfully. File size: {os.path.getsize(CHAT_HISTORY_FILE)} bytes")
        return True
    except Exception as e:
        print(f"Error saving chat history: {e}")
        return False

def load_ai_conversations():
    if os.path.exists(AI_CONVERSATIONS_FILE):
        try:
            with open(AI_CONVERSATIONS_FILE, 'r') as f:
                data = json.load(f)
                print(f"Successfully loaded AI conversations: {len(data)} conversations")
                return data
        except Exception as e:
            print(f"Error loading AI conversations: {e}")
    
    # Default empty conversations if file doesn't exist or is corrupted
    return {}

def save_ai_conversations(data):
    try:
        print(f"Saving AI conversations to {os.path.abspath(AI_CONVERSATIONS_FILE)}")
        with open(AI_CONVERSATIONS_FILE, 'w') as f:
            json.dump(data, f, indent=2)
        print(f"AI conversations saved successfully. File size: {os.path.getsize(AI_CONVERSATIONS_FILE)} bytes")
        return True
    except Exception as e:
        print(f"Error saving AI conversations: {e}")
        return False

@app.route('/api/resume/upload', methods=['POST'])
def upload_resume():
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file part"}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400

        if file:
            # Extract file extension (e.g., pdf, docx)
            ext = os.path.splitext(file.filename)[1].lower()
            filename = f"res{ext}"
            file_path = os.path.join(RESUMES_DIR, filename)

            # Save/overwrite the file
            file.save(file_path)
            print(f"Resume saved successfully as: {file_path}")

            return jsonify({
                "message": "Resume uploaded and saved as 'res' successfully",
                "filename": filename,
                "path": os.path.join("resumes", filename)
            })

    except Exception as e:
        print(f"Error uploading resume: {e}")
        return jsonify({"error": "Resume upload failed"}), 500

@app.route('/api/chat/save', methods=['POST'])
def save_chat():
    try:
        data = request.json
        chat_data = data.get('chatMessages', [])
        job_title = data.get('jobTitle', 'Untitled Job')
        
        # Load existing chat history
        chat_history = load_chat_history()
        
        # Create new chat entry with metadata
        new_chat = {
            "id": str(uuid.uuid4()),
            "jobTitle": job_title,
            "timestamp": datetime.now().isoformat(),
            "messages": chat_data
        }
        
        # Add to chat history
        chat_history.append(new_chat)
        
        # Save updated chat history
        success = save_chat_history(chat_history)
        
        if success:
            print(f"Chat history saved successfully for job: {job_title}")
            return jsonify({
                "message": "Chat history saved successfully",
                "chatId": new_chat["id"]
            })
        else:
            return jsonify({"error": "Failed to save chat history"}), 500
            
    except Exception as e:
        print(f"Error saving chat history: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/chat/history', methods=['GET'])
def get_chat_history():
    try:
        chat_history = load_chat_history()
        return jsonify(chat_history)
    except Exception as e:
        print(f"Error retrieving chat history: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/chat/respond', methods=['POST'])
def chat_respond():
    try:
        data = request.json
        messages = data.get("messages", [])

        # Construct OpenAI-compatible message format
        formatted_messages = [
            {"role": msg["role"], "content": msg["content"]}
            for msg in messages
        ]

        if not formatted_messages:
            return jsonify({"error": "No messages provided"}), 400

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Your name is ResumeCraft AI Agent, and you are a helpful assistant who improves job application documents."},
                *formatted_messages
            ]
        )

        reply = response.choices[0].message.content.strip()
        return jsonify({"reply": reply})

    except Exception as e:
        print(f"Error generating chat reply: {e}")
        return jsonify({"error": str(e)}), 500

        # Simple logic to infer target based on the last user message
        last_user_msg = next((msg["content"].lower() for msg in reversed(messages) if msg["role"] == "user"), "")
        if any(keyword in last_user_msg for keyword in ["resume", "cv", "bullet", "experience"]):
            target = "resume"
        elif any(keyword in last_user_msg for keyword in ["cover letter", "motivation", "introduction", "greeting"]):
            target = "coverLetter"
        else:
            target = "resume"  # default

        return jsonify({
            "reply": reply,
            "target": target
        })

    except Exception as e:
        print(f"Error generating chat reply: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/ai-conversation/save', methods=['POST'])
def save_ai_conversation():
    try:
        data = request.json
        conversation_id = data.get('conversationId')
        job_title = data.get('jobTitle', 'Untitled Job')
        messages = data.get('messages', [])
        
        if not conversation_id:
            conversation_id = str(uuid.uuid4())
        
        # Load existing AI conversations
        conversations = load_ai_conversations()
        
        # Create or update conversation
        conversations[conversation_id] = {
            "jobTitle": job_title,
            "lastUpdated": datetime.now().isoformat(),
            "messages": messages
        }
        
        # Save updated conversations
        success = save_ai_conversations(conversations)
        
        if success:
            print(f"AI conversation saved successfully for ID: {conversation_id}")
            return jsonify({
                "message": "Conversation saved successfully",
                "conversationId": conversation_id
            })
        else:
            return jsonify({"error": "Failed to save conversation"}), 500
            
    except Exception as e:
        print(f"Error saving AI conversation: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/ai-conversation/<conversation_id>', methods=['GET'])
def get_ai_conversation(conversation_id):
    try:
        conversations = load_ai_conversations()
        
        if conversation_id in conversations:
            return jsonify(conversations[conversation_id])
        else:
            return jsonify({"error": "Conversation not found"}), 404
    except Exception as e:
        print(f"Error retrieving AI conversation: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/ai-conversation/list', methods=['GET'])
def list_ai_conversations():
    try:
        conversations = load_ai_conversations()
        
        # Create a simplified list with just IDs, titles, and timestamps
        conversation_list = [
            {
                "id": conv_id,
                "jobTitle": data.get("jobTitle", "Untitled Job"),
                "lastUpdated": data.get("lastUpdated"),
                "messageCount": len(data.get("messages", [])) if "messages" in data else 0
            }
            for conv_id, data in conversations.items()
        ]
        
        return jsonify(conversation_list)
    except Exception as e:
        print(f"Error listing AI conversations: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/resume/generate', methods=['POST'])
def generate_resume():
    data = request.json
    job_title = data.get('jobTitle', '')
    job_description = data.get('jobDescription', '')

    try:
        user_data = load_user_data()
        education = user_data.get("education", [])
        work_experience = user_data.get("workExperiences", [])
        skills = user_data.get("skills", [])

        # -------- Resume Prompt --------
        prompt = f"""
You are a professional resume writer helping tailor resumes for a specific job description. Keep it recruiter-friendly and ATS-compliant. Use strong action verbs, quantified achievements, and align each bullet point with the provided job description.

The goal is to **rewrite only the WORK EXPERIENCE section** from the candidate's resume. Your output will be injected into an existing resume layout, so do NOT include any summary, contact info, education, or skills — just the updated WORK EXPERIENCE section in clean resume bullet format.

Candidate's Original Work Experience:
""" + "\n\n".join([
    f"{exp['position']} | {exp['company']} | {exp['startDate']} – {exp['endDate']}\n" +
    "\n".join([f"• {bullet}" for bullet in exp["bullets"]])
    for exp in work_experience
]) + f"""

Target Job Description:
{job_description}

Instructions:
- Focus on aligning the content with the sourcing, analytics, supplier governance, and software strategy focus in the job description.
- Preserve formatting: Return output using the same format as above (Job title | Company | Dates + 3–5 bullet points each).
- Do not add or remove job roles — just rewrite the bullet points to better match the job.
- Keep a concise, executive MBA-style tone.
"""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert resume writer skilled in ATS optimization."},
                {"role": "user", "content": prompt}
            ]
        )

        updated_experience = response.choices[0].message.content.strip()

        # Check for uploaded resume file
        resume_file = None
        for fname in os.listdir(RESUMES_DIR):
            if fname.startswith("res"):
                resume_file = os.path.join(RESUMES_DIR, fname)
                break

        if not resume_file:
            return jsonify({"error": "No uploaded resume file found. Please upload one named 'res'."}), 400

        now = datetime.now().isoformat()
        resume_id = str(uuid.uuid4())

        user_data["resumes"].append({
            "id": resume_id,
            "title": f"Updated Resume for {job_title}",
            "content": updated_experience,
            "createdAt": now
        })

        # -------- Cover Letter Prompt --------
        cover_prompt = f"""
Write a concise, personalized cover letter for the following job title and description.

Job Title: {job_title}
Job Description:
{job_description}

Candidate Background:
Skills: {", ".join(skills)}
Recent Experience Highlights:
""" + "\n".join([
    f"{exp['position']} at {exp['company']} – Key Impact: {exp['bullets'][0]}"
    for exp in work_experience[:2]
]) + """

Make it sound confident, polished, and tailored to the role.
"""

        cover_response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a skilled business communicator who writes concise, effective cover letters."},
                {"role": "user", "content": cover_prompt}
            ]
        )

        # -------- Format Cover Letter Header --------
        name = user_data.get("name", "[Your Name]")
        email = user_data.get("email", "[Your Email]")
        phone = user_data.get("phone", "[Your Phone Number]")
        date_str = datetime.now().strftime("%B %d, %Y")

        formatted_header = "\n".join([
            name,
            "[Your Address]",
            "[City, State, Zip]",
            email,
            phone,
            date_str
        ])

        # Replace the top 6 lines of the GPT-generated letter with our real info
        cover_letter_lines = cover_response.choices[0].message.content.strip().split("\n")
        body_start_index = 6 if len(cover_letter_lines) >= 6 else 0
        body_rest = "\n".join(cover_letter_lines[body_start_index:])
        # Remove any existing "Warm regards," block in the GPT body
        # Format your actual header
        name = user_data.get("name", "[Your Name]")
        email = user_data.get("email", "[Your Email]")
        phone = user_data.get("phone", "[Your Phone Number]")
        date_str = datetime.now().strftime("%B %d, %Y")

        header_lines = [
            name,
            "[Your Address]",
            "[City, State, Zip]",
            email,
            phone,
            date_str
        ]
        formatted_header = "\n".join(header_lines)

        # Get GPT's full output as lines
        gpt_lines = cover_response.choices[0].message.content.strip().split("\n")

        # Replace just the first 6 lines (header), keep the rest as-is
        body_rest = "\n".join(gpt_lines[6:]) if len(gpt_lines) > 6 else ""

        # Final combined cover letter text
        cover_letter_text = f"{formatted_header}\n{body_rest}".strip()
        if "[Your Name]" in cover_letter_text:
            cover_letter_text = cover_letter_text.replace("[Your Name]", name)


        # Save cover letter
        cover_letter_id = str(uuid.uuid4())
        user_data["coverLetters"].append({
            "id": cover_letter_id,
            "title": f"Cover Letter for {job_title}",
            "content": cover_letter_text,
            "createdAt": now
        })

        save_user_data(user_data)
        
        TEMP_FILE = os.path.join(DATA_DIR, "tmp", "generated_resume.json")
        os.makedirs(os.path.dirname(TEMP_FILE), exist_ok=True)

        with open(TEMP_FILE, "w") as f:
            json.dump({
                "resume": updated_experience,
                "coverLetter": cover_letter_text
            }, f)
        return jsonify({
            "resume": updated_experience,
            "coverLetter": cover_letter_text
        })

    except Exception as e:
        print(f"Error in resume generation: {e}")
        return jsonify({"error": "Failed to generate resume. Please try again."}), 500

@app.route('/api/document/download', methods=['POST'])
def download_document():
    try:
        data = request.json
        content = data.get('content', '')
        file_name = data.get('fileName', 'document')
        format_type = data.get('format', 'pdf')

        # Decide type based on whether content is provided
        if not content.strip():  # empty means it's resume
            user_data = load_user_data()
            if not user_data.get("resumes"):
                return jsonify({"error": "No resume found"}), 404
            content = user_data["resumes"][-1]["content"]

        if not content:
            return jsonify({"error": "No content available"}), 400

        if format_type == 'pdf':
            if content.strip().lower().startswith("dear") or "dear hiring manager" in content.lower():
                pdf_path = generate_cover_letter_pdf(content, file_name)
            else:
                pdf_path = generate_resume_pdf(content, file_name)

            if pdf_path and os.path.exists(pdf_path):
                return send_file(pdf_path, mimetype='application/pdf', as_attachment=True)
            else:
                return jsonify({"error": "Failed to generate PDF"}), 500

        elif format_type == 'docx':
            if content.strip().lower().startswith("dear") or "dear hiring manager" in content.lower():
                return generate_cover_letter_docx(content, file_name)
            else:
                return generate_resume_docx(content, file_name)

        else:
            return jsonify({"error": "Unsupported format"}), 400

    except Exception as e:
        print(f"Error in document download: {e}")
        return jsonify({"error": str(e)}), 500

def extract_additional_sections():
    additional_info = {}
    if os.path.exists(RESUMES_DIR):
        for fname in os.listdir(RESUMES_DIR):
            if fname.startswith("res") and fname.endswith(".docx"):
                doc_path = os.path.join(RESUMES_DIR, fname)
                doc = Document(doc_path)
                for table in doc.tables:
                    for row in table.rows:
                        if len(row.cells) >= 2:
                            header = row.cells[0].text.strip().rstrip(":")
                            data = row.cells[1].text.strip()
                            if header and data:
                                additional_info[header] = data.split('\n')
                break
    return additional_info

def generate_resume_pdf(content, file_name):
    with open(USER_DATA_FILE, "r") as f:
        user_data = json.load(f)

    additional_info = extract_additional_sections()
    extra_skills = additional_info.pop("Skills", [])

    output_path = os.path.join(DATA_DIR, f"{file_name}.pdf")
    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            leftMargin=0.5 * inch, rightMargin=0.5 * inch,
                            topMargin=0.5 * inch, bottomMargin=0.5 * inch)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Header', fontName='Times-Roman', fontSize=18, alignment=TA_CENTER, spaceAfter=10))
    styles.add(ParagraphStyle(name='Section', fontName='Times-Roman', fontSize=14, spaceBefore=12, spaceAfter=6, textColor=colors.HexColor('#2c3e50')))
    styles.add(ParagraphStyle(name='Body', fontName='Times-Roman', fontSize=11, spaceAfter=4))
    styles.add(ParagraphStyle(name='MyBullet', fontName='Times-Roman', fontSize=11, leftIndent=10, bulletIndent=5))

    elements = []

    elements.append(Paragraph(user_data["name"], styles["Header"]))
    contact = f"{user_data['email']} | {user_data['phone']} | {user_data['linkedin']}"
    elements.append(Paragraph(contact, styles["Body"]))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("Education", styles["Section"]))
    for edu in user_data.get("education", []):
        edu_summary = f"{edu['degree']} in {edu['field']} - {edu['institution']} ({edu['startDate']} to {edu['endDate']})"
        elements.append(Paragraph(edu_summary, styles["Body"]))
        for line in edu.get("description", "").split('\n'):
            if line.strip():
                elements.append(Paragraph(line.strip(), styles["MyBullet"]))

    elements.append(Paragraph("Experience", styles["Section"]))
    if not content or not isinstance(content, str):
        content = "Experience details not available."
    for block in content.split("\n\n"):
        lines = block.strip().split("\n")
        if lines:
            elements.append(Paragraph(lines[0], styles["Body"]))
            for bullet in lines[1:]:
                if bullet.strip():
                    elements.append(Paragraph(bullet.strip(), styles["MyBullet"]))

    elements.append(Paragraph("Skills", styles["Section"]))
    all_skills = user_data.get("skills", []) + extra_skills
    skills_text = ", ".join(all_skills)
    elements.append(Paragraph(skills_text, styles["Body"]))

    if additional_info:
        elements.append(Paragraph("Additional Information", styles["Section"]))
        for section, items in additional_info.items():
            elements.append(Paragraph(f"<b>{section}</b>", styles["Body"]))
            for item in items:
                if item.strip():
                    elements.append(Paragraph(f"- {item}", styles["MyBullet"]))

    try:
        doc.build(elements)
        return output_path
    except Exception as e:
        print(f"Error while building PDF: {e}")
        return None

def generate_cover_letter_pdf(content, file_name):
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT

    output_path = os.path.join(DATA_DIR, f"{file_name}.pdf")

    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            leftMargin=0.75 * inch, rightMargin=0.75 * inch,
                            topMargin=0.75 * inch, bottomMargin=0.75 * inch)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Letter', fontName='Times-Roman', fontSize=12, leading=16, alignment=TA_LEFT))

    elements = []

    for line in content.strip().split("\n"):
        if line.strip():
            elements.append(Paragraph(line.strip(), styles["Letter"]))
            elements.append(Spacer(1, 6))

    try:
        doc.build(elements)
        return output_path
    except Exception as e:
        print(f"Error building cover letter PDF: {e}")
        return None


def generate_resume_docx(content, file_name):
    with open(USER_DATA_FILE, "r") as f:
        user_data = json.load(f)

    additional_info = extract_additional_sections()
    extra_skills = additional_info.pop("Skills", [])

    doc = Document()
    doc.add_heading(user_data["name"], level=0)
    contact = f"{user_data['email']} | {user_data['phone']} | {user_data['linkedin']}"
    doc.add_paragraph(contact)

    doc.add_heading("Education", level=1)
    for edu in user_data.get("education", []):
        edu_summary = f"{edu['degree']} in {edu['field']} - {edu['institution']} ({edu['startDate']} to {edu['endDate']})"
        doc.add_paragraph(edu_summary, style='Normal')
        for line in edu.get("description", "").split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip(), style='ListBullet')

    doc.add_heading("Experience", level=1)
    if not content or not isinstance(content, str):
        content = "Experience details not available."
    for block in content.split("\n\n"):
        lines = block.strip().split("\n")
        if lines:
            doc.add_paragraph(lines[0], style='Normal')
            for bullet in lines[1:]:
                if bullet.strip():
                    doc.add_paragraph(bullet.strip(), style='ListBullet')

    doc.add_heading("Skills", level=1)
    all_skills = user_data.get("skills", []) + extra_skills
    skills_text = ", ".join(all_skills)
    doc.add_paragraph(skills_text)

    if additional_info:
        doc.add_heading("Additional Information", level=1)
        for section, items in additional_info.items():
            doc.add_paragraph(section + ":", style='Normal')
            for item in items:
                if item.strip():
                    doc.add_paragraph(item.strip(), style='ListBullet')

    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)

    return send_file(
        docx_io,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"{file_name}.docx"
    )

def generate_cover_letter_docx(content, file_name):
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Style setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for line in content.strip().split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph("")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"{file_name}.docx"
    )


def generate_experience_section(experiences):
    if not experiences:
        return "No previous work experience."
    
    result = ""
    for exp in experiences[:3]:  # Limit to 3 experiences for brevity
        company = exp.get('company', 'Company Name')
        position = exp.get('position', 'Position')
        description = exp.get('description', '')
        bullets = exp.get('bullets', [])
        
        result += f"### {position} at {company}\n"
        if description:
            result += f"{description}\n"
        
        if bullets:
            for bullet in bullets[:3]:  # Limit to 3 bullets
                result += f"- {bullet}\n"
        
        result += "\n"
    
    return result

def generate_education_section(education):
    if not education:
        return "No formal education listed."
    
    result = ""
    for edu in education[:2]:  # Limit to 2 education entries
        institution = edu.get('institution', 'University Name')
        degree = edu.get('degree', 'Degree')
        field = edu.get('field', 'Field of Study')
        
        result += f"### {degree} in {field}\n"
        result += f"{institution}\n\n"
    
    return result

def generate_experience_summary(experiences):
    if not experiences:
        return "various fields"
    
    positions = [exp.get('position', 'professional roles') for exp in experiences[:2]]
    return " and ".join(positions)

@app.route('/api/profile', methods=['GET'])
def get_profile():
    user_data = load_user_data()
    return jsonify(user_data)

@app.route('/api/profile/update', methods=['POST'])
def update_profile():
    try:
        data = request.json
        print(f"Received profile update request with data: {data}")
        
        user_data = load_user_data()
        
        # Update user profile with provided data
        for key, value in data.items():
            if key in user_data:
                user_data[key] = value
        
        # Save the updated user data
        success = save_user_data(user_data)
        if success:
            print("Profile updated successfully!")
            return jsonify({"message": "Profile updated successfully"})
        else:
            print("Failed to save user data")
            return jsonify({"error": "Failed to save profile data"}), 500
    except Exception as e:
        print(f"Error in update_profile: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/profile/education/<education_id>', methods=['PUT'])
def update_education_entry(education_id):
    try:
        data = request.json
        print(f"Received education update for ID {education_id}: {data}")
        
        user_data = load_user_data()
        
        # Find and update the education entry
        found = False
        for i, edu in enumerate(user_data.get('education', [])):
            if edu.get('id') == education_id:
                # Update the education entry with the new data
                user_data['education'][i] = {**edu, **data}
                found = True
                break
        
        if not found:
            print(f"Education entry not found: {education_id}")
            return jsonify({"error": "Education entry not found"}), 404
            
        success = save_user_data(user_data)
        if success:
            print(f"Education entry {education_id} updated successfully")
            return jsonify({"message": "Education entry updated successfully", "education": user_data['education'][i]})
        else:
            return jsonify({"error": "Failed to save education data"}), 500
    except Exception as e:
        print(f"Error updating education entry: {e}")
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    # Check if the data directory exists
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
        print(f"Created directory: {DATA_DIR}")
    
    # Initialize user data file if it doesn't exist
    if not os.path.exists(USER_DATA_FILE):
        initial_data = {
            "name": "John Doe",
            "email": "john@example.com",
            "phone": "(123) 456-7890",
            "linkedin": "linkedin.com/in/johndoe",
            "education": [],
            "workExperiences": [],
            "skills": ["JavaScript", "React", "Python"],
            "resumes": [],
            "coverLetters": []
        }
        save_user_data(initial_data)
        print(f"Created initial user data file at {os.path.abspath(USER_DATA_FILE)}")
    else:
        print(f"Using existing user data file at {os.path.abspath(USER_DATA_FILE)}")
    
    # Initialize chat history file if it doesn't exist
    if not os.path.exists(CHAT_HISTORY_FILE):
        save_chat_history([])
        print(f"Created empty chat history file at {os.path.abspath(CHAT_HISTORY_FILE)}")
    else:
        print(f"Using existing chat history file at {os.path.abspath(CHAT_HISTORY_FILE)}")
        
    # Initialize AI conversations file if it doesn't exist
    if not os.path.exists(AI_CONVERSATIONS_FILE):
        save_ai_conversations({})
        print(f"Created empty AI conversations file at {os.path.abspath(AI_CONVERSATIONS_FILE)}")
    else:
        print(f"Using existing AI conversations file at {os.path.abspath(AI_CONVERSATIONS_FILE)}")
    
    print("Checking required packages...")
    required_packages = ["markdown", "xhtml2pdf", "python-docx"]
    missing_packages = []
    
    # Import each package to check if it's installed
    for package in required_packages:
        try:
            __import__(package)
            print(f"✓ {package} is installed")
        except ImportError:
            print(f"✗ {package} is not installed")
            missing_packages.append(package)
    
    if missing_packages:
        print("\nWarning: Some required packages are missing.")
        print("Please install them using the following command:")
        print(f"pip install {' '.join(missing_packages)}")
    else:
        print("\nAll required packages are installed!")
    
    # Run the Flask app
    print(f"\nStarting Flask server on http://localhost:5000")
    app.run(debug=True, port=5000)
